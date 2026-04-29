import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"c:\Users\hp\Downloads\sars_goal1")
TEX_PATH = ROOT / "Research_template" / "main.tex"
DOCX_PATH = ROOT / "Henweis_Format.docx"
BACKUP_PATH = ROOT / "Henweis_Format.backup.docx"


def extract_block(text: str, start: str, end: str) -> str:
    m = re.search(re.escape(start) + r"(.*?)" + re.escape(end), text, re.S)
    return m.group(1).strip() if m else ""


def strip_latex(s: str) -> str:
    # Preserve equation content as plain lines.
    s = re.sub(r"\\begin\{equation\}(.*?)\\end\{equation\}", lambda m: "\nEquation: " + m.group(1).strip() + "\n", s, flags=re.S)

    # Drop figure/table/algorithm environments (captions are usually descriptive elsewhere).
    s = re.sub(r"\\begin\{figure\*?\}.*?\\end\{figure\*?\}", "\n", s, flags=re.S)
    s = re.sub(r"\\begin\{table\*?\}.*?\\end\{table\*?\}", "\n", s, flags=re.S)
    s = re.sub(r"\\begin\{algorithm\}.*?\\end\{algorithm\}", "\n", s, flags=re.S)

    # Remove citations and refs.
    s = re.sub(r"\\cite\{[^}]*\}", "", s)
    s = re.sub(r"\\ref\{[^}]*\}", "", s)
    s = re.sub(r"\\label\{[^}]*\}", "", s)

    # Unwrap common formatting commands.
    s = re.sub(r"\\texttt\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\textbf\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\emph\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\\max\\left\((.*?)\\right\)", r"max(\1)", s)

    # Remove remaining latex commands but keep line breaks sensible.
    s = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?(\{[^}]*\})?", "", s)

    # Normalize escaped chars.
    s = s.replace("\\%", "%").replace("\\_", "_").replace("~", " ")
    s = s.replace("$", "")

    # Cleanup whitespace.
    s = re.sub(r"\n{3,}", "\n\n", s)
    s = re.sub(r"[ \t]{2,}", " ", s)
    return s.strip()


def add_paragraphs(doc: Document, text: str) -> None:
    for para in [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]:
        doc.add_paragraph(para)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_section_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True


tex = TEX_PATH.read_text(encoding="utf-8")

# Extract metadata.
title_match = re.search(r"\\title\{([^}]*)\}", tex)
title = title_match.group(1).strip() if title_match else "SARS Paper"

authors = re.findall(r"\\author\[[^\]]+\]\{([^}]*)\}", tex)
affils = re.findall(r"\\affil\[[^\]]*\]\{([^}]*)\}", tex)

abstract = extract_block(tex, r"\begin{abstract}", r"\end{abstract}")
keywords_match = re.search(r"\\noindent\\textbf\{Keywords:\}\s*(.*)", tex)
keywords = keywords_match.group(1).strip() if keywords_match else ""

# Extract sections.
section_matches = list(re.finditer(r"\\section\{([^}]*)\}", tex))
sections = []
for i, m in enumerate(section_matches):
    name = m.group(1).strip()
    start = m.end()
    end = section_matches[i + 1].start() if i + 1 < len(section_matches) else tex.find("\\section*{Acknowledgments}")
    if end == -1:
        end = len(tex)
    content = tex[start:end].strip()
    sections.append((name, content))

ack = extract_block(tex, r"\section*{Acknowledgments}", r"\printbibliography")

# Backup old file and rebuild using the same DOCX template file.
if DOCX_PATH.exists():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)

doc = Document(str(DOCX_PATH))
clear_document_body(doc)

p = doc.add_paragraph()
r = p.add_run(title)
r.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

if authors:
    p = doc.add_paragraph(", ".join(authors))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

if affils:
    for a in affils:
        p = doc.add_paragraph(strip_latex(a))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_section_title(doc, "Abstract")
add_paragraphs(doc, strip_latex(abstract))

if keywords:
    add_section_title(doc, "Keywords")
    add_paragraphs(doc, strip_latex(keywords))

for name, content in sections:
    add_section_title(doc, strip_latex(name))
    cleaned = strip_latex(content)
    if cleaned:
        add_paragraphs(doc, cleaned)

if ack:
    add_section_title(doc, "Acknowledgments")
    add_paragraphs(doc, strip_latex(ack))

add_section_title(doc, "References")
doc.add_paragraph("Please compile and format references from sample.bib according to the Henweis journal style.")

doc.save(str(DOCX_PATH))
print(f"Updated: {DOCX_PATH}")
print(f"Backup: {BACKUP_PATH}")
